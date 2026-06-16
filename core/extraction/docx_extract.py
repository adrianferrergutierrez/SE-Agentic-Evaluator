"""
core/extraction/docx_extract.py
================================
Tool: extract text, tables and images from a DOCX document.

Converts a Word document into structured Markdown and extracts
embedded images to a separate directory.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Dict, Any

logger = logging.getLogger(__name__)

MIN_API_IMAGE_DIMENSION = 11  # DashScope requires width/height > 10
MIN_IMAGE_BLOB_SIZE_BYTES = 128  # Fast-path to skip likely spacer pixels/icons


def extract_docx(input: str, output_dir: str) -> Dict[str, Any]:
    """Extract content from a DOCX file.

    Parameters
    ----------
    input:
        Path to the DOCX file.
    output_dir:
        Directory to write contents.md and img/ subdirectory.

    Returns
    -------
    dict
        Keys: ``contents_md`` (path to generated Markdown), ``images`` (count).
    """
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    try:
        from PIL import Image
        import io
        has_pil = True
    except ImportError:
        has_pil = False
        logger.warning("PIL not available. Tiny-image filtering will use blob-size fallback only.")

    input_path = Path(input)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    (output_path / "img").mkdir(exist_ok=True)

    doc = docx.Document(str(input_path))
    lines = []
    img_count = 0

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        if "Title" in style:
            lines.append(f"# {text}")
        elif "Heading 1" in style:
            lines.append(f"## {text}")
        elif "Heading 2" in style:
            lines.append(f"### {text}")
        elif "Heading 3" in style:
            lines.append(f"#### {text}")
        elif "List" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for i, table in enumerate(doc.tables):
        lines.append(f"\n### Tabla {i+1}\n")
        header = [cell.text for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.replace("|", "\\|") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
            img_name = f"img_{i}.{ext}"
            img_path = output_path / "img" / img_name
            try:
                blob = rel.target_part.blob

                # Skip tiny decorative images (common 1x1 spacers) before writing markdown refs.
                if len(blob) < MIN_IMAGE_BLOB_SIZE_BYTES:
                    logger.info("Skipping tiny embedded image by size: %s (%d bytes)", img_name, len(blob))
                    continue

                if has_pil:
                    try:
                        with Image.open(io.BytesIO(blob)) as im:
                            if im.width < MIN_API_IMAGE_DIMENSION or im.height < MIN_API_IMAGE_DIMENSION:
                                logger.info(
                                    "Skipping tiny embedded image by dimensions: %s (%dx%d)",
                                    img_name,
                                    im.width,
                                    im.height,
                                )
                                continue
                    except Exception as e:
                        logger.warning("Could not inspect image %s dimensions: %s", img_name, e)

                with open(img_path, "wb") as f:
                    f.write(blob)
                lines.append(f'\n![Imagen {i+1}](img/{img_name})\n')
                img_count += 1
            except Exception:
                pass

    md_content = "\n\n".join(lines)
    md_path = output_path / "contents.md"
    md_path.write_text(md_content, encoding="utf-8")

    logger.info("Extracted %d chars, %d images from %s", len(md_content), img_count, input)
    return {"contents_md": str(md_path), "images": img_count}
